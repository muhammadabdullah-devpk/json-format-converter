import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # top, bottom, left, right are in dxa (1 pt = 20 dxa)
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_proposal():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles config
    styles = doc.styles
    
    # Configure Normal style
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(45, 55, 72)  # Dark Gray
    
    # Colors
    NAVY = RGBColor(26, 54, 93)
    SLATE = RGBColor(43, 108, 176)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("SOFTWARE DEVELOPMENT PROPOSAL")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Daikibo Manufacturing Status Dashboard")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = SLATE
    subtitle.paragraph_format.space_after = Pt(36)
    
    # Metadata Box
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run("Client: Daikibo Steel & Manufacturing\nPrepared By: Advanced Software Engineering Team\nDate: June 30, 2026")
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(113, 128, 150)
    meta.paragraph_format.space_after = Pt(24)
    
    # Divider Line
    p_div = doc.add_paragraph()
    p_div_run = p_div.add_run("─" * 60)
    p_div_run.font.color.rgb = RGBColor(226, 232, 240)
    p_div.paragraph_format.space_after = Pt(24)
    
    # Add headings helper
    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        return p
        
    def add_bullet_point(p, bold_prefix, text):
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        p.style = 'List Bullet'
        r_bold = p.add_run(bold_prefix)
        r_bold.font.bold = True
        p.add_run(text)

    # 1. Executive Overview
    add_section_heading("1. Executive Overview")
    p1 = doc.add_paragraph(
        "Daikibo currently collects telemetry data from its manufacturing processes across multiple sites. "
        "While offline analysis provides valuable historical insights, it lacks the immediacy required "
        "to prevent downtime and respond rapidly to equipment failures."
    )
    p1.paragraph_format.space_after = Pt(10)
    p1.paragraph_format.line_spacing = 1.15
    
    p2 = doc.add_paragraph(
        "This proposal outlines the development of the Daikibo Manufacturing Status Dashboard, a secure, "
        "real-time single-page web application. The dashboard will monitor the health status of all 9 machines "
        "operating across each of Daikibo's 4 factories (totaling 36 monitored devices). By rendering telemetry "
        "data into an interactive, intuitive interface within the company's intranet, managers and operators "
        "can track live health indicators, view status histories, and receive immediate alerts for system anomalies."
    )
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.line_spacing = 1.15

    # 2. Project Scope
    add_section_heading("2. Project Scope & Functional Requirements")
    p3 = doc.add_paragraph(
        "The application will be developed as a secure, responsive single-page dashboard tailored to Daikibo's "
        "internal infrastructure. Core features of the system include:"
    )
    p3.paragraph_format.space_after = Pt(6)
    
    # Bullet points
    add_bullet_point(doc.add_paragraph(), "Single-Page Interface: ", "A unified screen presenting a clean list of all 36 monitored devices grouped by their respective factories.")
    add_bullet_point(doc.add_paragraph(), "Factory-Level Collapse/Expand: ", "Users can collapse or expand factory panels (Tokyo, Osaka, Nagoya, Fukuoka) to optimize screen real estate.")
    add_bullet_point(doc.add_paragraph(), "Device-Level Collapse/Expand: ", "Expanding a specific device reveals its detailed telemetry history, status codes, and a graphical timeline of past performance.")
    add_bullet_point(doc.add_paragraph(), "Intranet-Only Access: ", "The dashboard will be hosted internally, ensuring that data never leaves the secure boundaries of Daikibo’s corporate intranet.")
    add_bullet_point(doc.add_paragraph(), "Single Sign-On (SSO) Integration: ", "User authentication will sync with the client's internal Active Directory/LDAP server, enabling standard company accounts.")
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(8)

    # 3. Technical Architecture
    add_section_heading("3. Technical Architecture & Tech Stack")
    p4 = doc.add_paragraph(
        "To ensure security, performance, and long-term maintainability, we propose the following modern technology stack:"
    )
    p4.paragraph_format.space_after = Pt(6)
    
    add_bullet_point(doc.add_paragraph(), "Frontend UI: ", "React.js / TypeScript styled with a premium CSS grid/flexbox system for responsiveness on tablet and desktop terminals.")
    add_bullet_point(doc.add_paragraph(), "Data Visualization: ", "Chart.js or Recharts for lightweight, high-performance telemetry history rendering.")
    add_bullet_point(doc.add_paragraph(), "Backend API Service: ", "Python (FastAPI) or Node.js (Express), designed to serve cached telemetry states and query the database efficiently.")
    add_bullet_point(doc.add_paragraph(), "Database Integration: ", "PostgreSQL or TimescaleDB for time-series telemetry storage.")
    add_bullet_point(doc.add_paragraph(), "Security & Authentication: ", "Integration with Active Directory / LDAP using an OAuth2 proxy. SSL/TLS encryption for all internal network traffic.")

    # 4. Estimation & Resource Breakdown
    add_section_heading("4. Estimation & Resource Breakdown")
    p5 = doc.add_paragraph(
        "This estimate covers all phases of the software development lifecycle (SDLC), including requirements gathering, architecture, development, testing, and deployment."
    )
    p5.paragraph_format.space_after = Pt(12)
    
    # Table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Define widths
    widths = [Inches(3.5), Inches(2.0), Inches(1.0)]
    
    headers = ["Phase & Description", "Details", "Est. Hours"]
    for i, name in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = name
        set_cell_background(cell, "1A365D")  # Navy hex
        # Text style for header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    
    data = [
        ["Phase 1: Architecture & Auth Integration", "System design & AD/LDAP sync prototype", "40 hrs"],
        ["Phase 2: Backend API & Telemetry Pipeline", "RESTful APIs & Database connections", "60 hrs"],
        ["Phase 3: Frontend Dashboard Development", "Single-page UI, collapsible grid, charts", "80 hrs"],
        ["Phase 4: QA & Automated Testing", "Unit/Integration tests, auth simulation", "40 hrs"],
        ["Phase 5: Intranet Deployment & Handover", "Production deploy, training, docs", "30 hrs"],
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            # Padding and sizing
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")  # Alternating gray
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if col_idx == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    
    # Total row
    row_total = table.rows[-1]
    # Let's bold the last row
    table.cell(5, 0).text = "Total Estimated Effort"
    table.cell(5, 1).text = ""
    table.cell(5, 2).text = "250 hrs"
    for i in range(3):
        cell = table.cell(5, i)
        set_cell_background(cell, "EDF2F7")
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = NAVY
                
    p_spacer2 = doc.add_paragraph()
    p_spacer2.paragraph_format.space_before = Pt(12)

    # 5. Project Timeline
    add_section_heading("5. Project Timeline & Milestones")
    p6 = doc.add_paragraph(
        "The proposed timeline spans a 6-week development cycle based on a dedicated team of engineers:"
    )
    p6.paragraph_format.space_after = Pt(6)
    
    add_bullet_point(doc.add_paragraph(), "Milestone 1 (End of Week 1): ", "Architecture sign-off & Successful SSO/AD authentication prototype.")
    add_bullet_point(doc.add_paragraph(), "Milestone 2 (End of Week 3): ", "Backend API completion and functional telemetry database integration.")
    add_bullet_point(doc.add_paragraph(), "Milestone 3 (End of Week 4): ", "Frontend UI completed with interactive factory/device collapsible views.")
    add_bullet_point(doc.add_paragraph(), "Milestone 4 (End of Week 5): ", "Testing cycle complete with zero unresolved high-severity bugs.")
    add_bullet_point(doc.add_paragraph(), "Milestone 5 (End of Week 6): ", "Successful production deploy on Daikibo intranet and final handover.")

    # 6. Support
    add_section_heading("6. Continuous Support & Maintenance")
    p7 = doc.add_paragraph(
        "Post-deployment, Daikibo can rely on our continuous support program to ensure the system remains secure, performant, and up-to-date:"
    )
    p7.paragraph_format.space_after = Pt(6)
    
    add_bullet_point(doc.add_paragraph(), "Warranty Period: ", "30 days of complimentary support covering all critical bug fixes and deployment stability issues.")
    add_bullet_point(doc.add_paragraph(), "SLA Support Packages: ", "Tiered maintenance contracts including Standard (9/5 support with 4-hour response) and Premium (24/7 support with 1-hour response).")
    add_bullet_point(doc.add_paragraph(), "Continuous Updates: ", "Regular security patching, framework upgrades, and minor feature adjustments included.")
    
    # Save the document
    doc.save("Daikibo_Development_Proposal.docx")
    print("Successfully generated Daikibo_Development_Proposal.docx")

if __name__ == "__main__":
    create_proposal()
